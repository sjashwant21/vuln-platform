"""
DOCX renderer — python-docx report generation.

Generates a structured Word document from ReportData.
Charts are embedded as PNG images (python-docx requires raster).

Design:
  _DocxBuilder is a thin wrapper around python-docx that adds
  organisation-specific styling (corporate fonts, colour palette,
  table styles) via a shared document template approach.

  Every section method returns self for method chaining readability.
"""
from __future__ import annotations

import asyncio
import io
from datetime import datetime
from typing import Any

import structlog

from app.domain.models.report import (
    ReportData, ReportType, SeverityLevel, VulnDetail, AssetSummary,
)

logger = structlog.get_logger(__name__)

# RGB tuples for python-docx RGBColor
_RGB = {
    "critical": (0xDC, 0x26, 0x26),
    "high":     (0xEA, 0x58, 0x0C),
    "medium":   (0xD9, 0x77, 0x06),
    "low":      (0x25, 0x63, 0xEB),
    "ok":       (0x16, 0xA3, 0x4A),
    "accent":   (0x4F, 0x46, 0xE5),
    "text":     (0x11, 0x18, 0x27),
    "dim":      (0x6B, 0x72, 0x80),
    "white":    (0xFF, 0xFF, 0xFF),
    "bg":       (0xF9, 0xFA, 0xFB),
}


class DocxRenderer:
    """Renders ReportData to DOCX bytes via python-docx."""

    async def render(
        self,
        report_data: ReportData,
        charts_png:  dict[str, bytes],   # name → PNG bytes
    ) -> bytes:
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None, self._build_docx, report_data, charts_png
        )
        logger.info(
            "docx_rendered",
            report_type=report_data.report_type.value,
            size_kb=round(len(result) / 1024, 1),
        )
        return result

    def _build_docx(self, data: ReportData, charts: dict[str, bytes]) -> bytes:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Page margins ─────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # ── Cover page ────────────────────────────────────────
        self._add_cover(doc, data)
        doc.add_page_break()

        # ── Table of contents placeholder ─────────────────────
        self._add_heading(doc, "Contents", level=1)
        p = doc.add_paragraph()
        p.add_run("Executive Summary  ·  Risk Metrics  ·  Asset Inventory  ·  "
                  "Vulnerability Details  ·  Risk Trend  ·  Recommendations").italic = True
        doc.add_paragraph()

        # ── Sections by report type ───────────────────────────
        self._add_executive_section(doc, data, charts)

        if data.report_type in (ReportType.TECHNICAL, ReportType.VULNERABILITY):
            self._add_asset_section(doc, data, charts)
            self._add_vuln_section(doc, data, charts)

        if data.report_type == ReportType.COMPLIANCE:
            self._add_compliance_section(doc, data, charts)

        if data.risk_trend:
            self._add_trend_section(doc, data, charts)

        if data.ai_recommendations:
            self._add_recommendations_section(doc, data)

        # ── Footer note ───────────────────────────────────────
        self._add_footer(doc, data)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ── Section builders ──────────────────────────────────────

    def _add_cover(self, doc: Any, data: ReportData) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(data.org_name.upper())
        run.font.size = Pt(10)
        run.font.color.rgb = self._rgb("dim")
        run.font.bold = True

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run(f"{data.report_type.value.replace('_', ' ').title()} Report")
        run.font.size = Pt(11)
        run.font.color.rgb = self._rgb("accent")
        run.font.bold = True

        p = doc.add_paragraph()
        run = p.add_run("Security Assessment Report")
        run.font.size = Pt(26)
        run.font.color.rgb = self._rgb("accent")
        run.font.bold = True

        doc.add_paragraph()

        meta = [
            ("Generated", data.generated_at.strftime("%B %d, %Y at %H:%M UTC")),
            ("Period",    f"{data.period_start.strftime('%b %d')} – {data.period_end.strftime('%b %d, %Y')}"),
            ("Prepared by", data.generated_by),
            ("Report ID", data.report_id),
        ]
        for label, value in meta:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.font.bold = True
            run.font.size = Pt(9)
            p.add_run(value).font.size = Pt(9)

        doc.add_paragraph()
        p = doc.add_paragraph("⚠  CONFIDENTIAL — FOR INTERNAL USE ONLY  ⚠")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = self._rgb("dim")
            run.font.size = Pt(8)

    def _add_executive_section(self, doc: Any, data: ReportData, charts: dict) -> None:
        from docx.shared import Pt, Inches

        self._add_heading(doc, "Executive Summary", level=1)
        doc.add_paragraph(data.executive_summary)

        # Health score chart
        if "health_gauge" in charts:
            self._add_png(doc, charts["health_gauge"], width=Inches(3.5),
                          caption="Security Health Score")

        # Severity donut
        if "severity_donut" in charts:
            self._add_png(doc, charts["severity_donut"], width=Inches(3.5),
                          caption="Vulnerability Distribution")

        # Metrics table
        self._add_heading(doc, "Key Metrics", level=2)
        metrics = [
            ("Security Score",  str(data.health_score), "accent"),
            ("Critical Vulns",  str(data.severity_distribution.critical), "critical"),
            ("High Vulns",      str(data.severity_distribution.high),     "high"),
            ("Medium Vulns",    str(data.severity_distribution.medium),   "medium"),
            ("Low Vulns",       str(data.severity_distribution.low),      "low"),
            ("Assets Scanned",  str(data.total_assets),                   "accent"),
            ("Open Vulns",      str(data.open_vulns),                     "critical"),
            ("Resolved",        str(data.resolved_vulns),                 "ok"),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        self._set_cell_text(hdr[0], "Metric", bold=True)
        self._set_cell_text(hdr[1], "Value",  bold=True)
        for label, value, color in metrics:
            row  = table.add_row()
            self._set_cell_text(row.cells[0], label)
            self._set_cell_text(row.cells[1], value, color=color)
        doc.add_paragraph()

    def _add_asset_section(self, doc: Any, data: ReportData, charts: dict) -> None:
        from docx.shared import Inches, Pt

        doc.add_page_break()
        self._add_heading(doc, "Asset Inventory", level=1)

        if "asset_risk" in charts:
            self._add_png(doc, charts["asset_risk"], width=Inches(6),
                          caption="Asset Risk Scores")

        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = ["Host", "IP", "Type", "Criticality", "Ports", "C", "H", "Risk"]
        for i, h in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[i], h, bold=True)

        for asset in data.assets:
            row = table.add_row()
            self._set_cell_text(row.cells[0], asset.display_name[:20])
            self._set_cell_text(row.cells[1], asset.ip_address, mono=True)
            self._set_cell_text(row.cells[2], asset.asset_type)
            self._set_cell_text(row.cells[3], asset.criticality,
                                color=asset.criticality if asset.criticality in _RGB else "text")
            self._set_cell_text(row.cells[4], str(asset.open_ports))
            self._set_cell_text(row.cells[5], str(asset.vuln_critical),
                                color="critical" if asset.vuln_critical > 0 else "text")
            self._set_cell_text(row.cells[6], str(asset.vuln_high),
                                color="high" if asset.vuln_high > 0 else "text")
            self._set_cell_text(row.cells[7], f"{asset.risk_score:.1f}")
        doc.add_paragraph()

    def _add_vuln_section(self, doc: Any, data: ReportData, charts: dict) -> None:
        from docx.shared import Inches

        doc.add_page_break()
        self._add_heading(doc, "Vulnerability Details", level=1)

        if "asset_stacked" in charts:
            self._add_png(doc, charts["asset_stacked"], width=Inches(6),
                          caption="Vulnerabilities by Asset")

        open_vulns = [v for v in data.vulns if v.status == "open"]
        open_vulns.sort(key=lambda v: v.cvss_score or 0, reverse=True)

        for severity in ["critical", "high", "medium", "low"]:
            sev_vulns = [v for v in open_vulns if v.severity.value == severity]
            if not sev_vulns:
                continue

            self._add_heading(doc, f"{severity.title()} Vulnerabilities ({len(sev_vulns)})", level=2)

            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            for i, h in enumerate(["CVE", "Title", "CVSS", "Asset", "Age (d)", "Exploit"]):
                self._set_cell_text(table.rows[0].cells[i], h, bold=True)

            for v in sev_vulns:
                row = table.add_row()
                self._set_cell_text(row.cells[0], v.cve_id or "—", mono=True)
                self._set_cell_text(row.cells[1], v.title[:45])
                self._set_cell_text(row.cells[2], f"{v.cvss_score:.1f}" if v.cvss_score else "?",
                                    color=severity)
                self._set_cell_text(row.cells[3], v.asset_name[:18])
                self._set_cell_text(row.cells[4], str(v.age_days),
                                    color="critical" if v.age_days > 30 else "text")
                self._set_cell_text(row.cells[5], "⚠ YES" if v.has_exploit else "No",
                                    color="critical" if v.has_exploit else "text")
            doc.add_paragraph()

    def _add_compliance_section(self, doc: Any, data: ReportData, charts: dict) -> None:
        from docx.shared import Inches

        doc.add_page_break()
        self._add_heading(doc, "Compliance Assessment", level=1)

        if "compliance" in charts:
            self._add_png(doc, charts["compliance"], width=Inches(5.5),
                          caption="Compliance by Framework")

        for fw in data.compliance:
            self._add_heading(doc, f"{fw.framework} — {fw.compliance_pct}% Compliant", level=2)

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for i, h in enumerate(["Control", "Title", "Status", "Severity"]):
                self._set_cell_text(table.rows[0].cells[i], h, bold=True)

            for ctrl in fw.controls:
                row = table.add_row()
                self._set_cell_text(row.cells[0], ctrl.control_id, mono=True)
                self._set_cell_text(row.cells[1], ctrl.title[:45])
                color = "ok" if ctrl.status == "compliant" else \
                        "critical" if ctrl.status == "non_compliant" else "medium"
                self._set_cell_text(row.cells[2], ctrl.status.replace("_", " ").title(),
                                    color=color)
                self._set_cell_text(row.cells[3], ctrl.severity, color=ctrl.severity)
            doc.add_paragraph()

    def _add_trend_section(self, doc: Any, data: ReportData, charts: dict) -> None:
        from docx.shared import Inches

        doc.add_page_break()
        self._add_heading(doc, "Risk Trend Analysis", level=1)

        if "risk_trend" in charts:
            self._add_png(doc, charts["risk_trend"], width=Inches(6.5),
                          caption="Risk Trend (Last 30 Days)")

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        for i, h in enumerate(["Date", "Critical", "High", "Medium", "Low", "Score"]):
            self._set_cell_text(table.rows[0].cells[i], h, bold=True)
        for p in data.risk_trend:
            row = table.add_row()
            self._set_cell_text(row.cells[0], p.date.strftime("%Y-%m-%d"))
            self._set_cell_text(row.cells[1], str(p.critical),
                                color="critical" if p.critical > 0 else "text")
            self._set_cell_text(row.cells[2], str(p.high))
            self._set_cell_text(row.cells[3], str(p.medium))
            self._set_cell_text(row.cells[4], str(p.low))
            self._set_cell_text(row.cells[5], str(p.score))
        doc.add_paragraph()

    def _add_recommendations_section(self, doc: Any, data: ReportData) -> None:
        doc.add_page_break()
        self._add_heading(doc, "AI-Powered Recommendations", level=1)
        for rec in data.ai_recommendations:
            self._add_heading(doc, f"{rec.priority}. {rec.title}", level=2)
            doc.add_paragraph(rec.description)
            p = doc.add_paragraph()
            p.add_run("Effort: ").bold = True
            p.add_run(rec.effort.replace("_", " ").title())
            p2 = doc.add_paragraph()
            p2.add_run("Expected impact: ").bold = True
            p2.add_run(rec.impact)
            doc.add_paragraph()

    def _add_footer(self, doc: Any, data: ReportData) -> None:
        from docx.shared import Pt
        doc.add_paragraph()
        p = doc.add_paragraph(
            f"Generated {data.generated_at.strftime('%Y-%m-%d')} | "
            f"Report ID: {data.report_id} | CONFIDENTIAL"
        )
        for run in p.runs:
            run.font.size = Pt(7)
            run.font.color.rgb = self._rgb("dim")

    # ── Helpers ────────────────────────────────────────────────

    def _add_heading(self, doc: Any, text: str, level: int = 1) -> Any:
        from docx.shared import Pt, RGBColor
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = self._rgb("accent" if level == 1 else "text")
        return heading

    def _set_cell_text(
        self,
        cell: Any,
        text: str,
        bold: bool = False,
        color: str = "text",
        mono: bool = False,
    ) -> None:
        from docx.shared import Pt, RGBColor
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.bold = bold
        run.font.size = Pt(8)
        if color in _RGB:
            run.font.color.rgb = self._rgb(color)
        if mono:
            run.font.name = "Courier New"

    def _add_png(
        self,
        doc: Any,
        png_bytes: bytes,
        width: Any = None,
        caption: str = "",
    ) -> None:
        from docx.shared import Inches, Pt
        buf = io.BytesIO(png_bytes)
        doc.add_picture(buf, width=width)
        if caption:
            p = doc.add_paragraph(caption)
            for run in p.runs:
                run.font.size   = Pt(8)
                run.font.italic = True
                run.font.color.rgb = self._rgb("dim")

    @staticmethod
    def _rgb(key: str) -> Any:
        from docx.shared import RGBColor
        r, g, b = _RGB.get(key, _RGB["text"])
        return RGBColor(r, g, b)
